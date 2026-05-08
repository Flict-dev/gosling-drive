from __future__ import annotations

from pathlib import Path
from tempfile import NamedTemporaryFile
from zipfile import ZIP_DEFLATED, ZipFile

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "output"
SCREEN = OUT / "playwright"
ASSETS = ROOT / "report_assets"
DOCX_PATH = OUT / "КР_Жданов_Гослинг_Drive_отчет.docx"


figure_counter = 0
table_counter = 0
listing_counter = 0


def set_run_font(run, size: float = 14, bold: bool | None = None, italic: bool | None = None):
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:ascii"), "Times New Roman")
    r_fonts.set(qn("w:hAnsi"), "Times New Roman")
    r_fonts.set(qn("w:cs"), "Times New Roman")
    r_fonts.set(qn("w:eastAsia"), "Times New Roman")
    return run


def set_paragraph_base(paragraph, *, first_line: bool = True, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    fmt = paragraph.paragraph_format
    fmt.alignment = align
    fmt.first_line_indent = Cm(1.25) if first_line else Cm(0)
    fmt.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    for run in paragraph.runs:
        set_run_font(run)
    return paragraph


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    set_run_font(run)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def configure_document(doc: Document):
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(1.0)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.header_distance = Cm(1.25)
    section.footer_distance = Cm(1.25)
    section.different_first_page_header_footer = True
    add_page_number(section.footer.paragraphs[0])

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(14)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    pf = normal.paragraph_format
    pf.first_line_indent = Cm(1.25)
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)

    h1 = styles["Heading 1"]
    h1.font.name = "Times New Roman"
    h1.font.size = Pt(14)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor(0, 0, 0)
    h1._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    h1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h1.paragraph_format.first_line_indent = Cm(0)
    h1.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    h1.paragraph_format.space_before = Pt(0)
    h1.paragraph_format.space_after = Pt(0)
    h1.paragraph_format.keep_with_next = True

    h2 = styles["Heading 2"]
    h2.font.name = "Times New Roman"
    h2.font.size = Pt(14)
    h2.font.bold = True
    h2.font.color.rgb = RGBColor(0, 0, 0)
    h2._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    h2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    h2.paragraph_format.first_line_indent = Cm(1.25)
    h2.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    h2.paragraph_format.space_before = Pt(0)
    h2.paragraph_format.space_after = Pt(0)
    h2.paragraph_format.keep_with_next = True

    for style_name in ["List Bullet", "List Number"]:
        style = styles[style_name]
        style.font.name = "Times New Roman"
        style.font.size = Pt(14)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.paragraph_format.left_indent = Cm(1.25)
        style.paragraph_format.first_line_indent = Cm(0)
        style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        style.paragraph_format.space_after = Pt(0)


def paragraph(doc: Document, text: str = "", *, first_line: bool = True, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph()
    set_paragraph_base(p, first_line=first_line, align=align)
    if text:
        run = p.add_run(text)
        set_run_font(run)
    return p


def section_heading(doc: Document, text: str, *, page_break: bool = True):
    if page_break and len(doc.paragraphs) > 0:
        doc.add_page_break()
    p = doc.add_paragraph(text.upper(), style="Heading 1")
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    for run in p.runs:
        set_run_font(run, bold=True)
    return p


def subsection(doc: Document, text: str):
    p = doc.add_paragraph(text, style="Heading 2")
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.first_line_indent = Cm(1.25)
    for run in p.runs:
        set_run_font(run, bold=True)
    return p


def bullet(doc: Document, text: str):
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.left_indent = Cm(1.25)
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.paragraph_format.space_after = Pt(0)
    for run in p.runs:
        set_run_font(run)
    return p


def numbered(doc: Document, text: str):
    p = doc.add_paragraph(text, style="List Number")
    p.paragraph_format.left_indent = Cm(1.25)
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.paragraph_format.space_after = Pt(0)
    for run in p.runs:
        set_run_font(run)
    return p


def set_cell_margins(cell, top=120, start=120, bottom=120, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def shade_cell(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text: str, *, bold: bool = False, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_margins(cell)
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    set_run_font(run, size=12, bold=bold)


def add_table_caption(doc: Document, title: str):
    global table_counter
    table_counter += 1
    p = paragraph(doc, f"Таблица {table_counter} - {title}", first_line=False, align=WD_ALIGN_PARAGRAPH.LEFT)
    p.paragraph_format.keep_with_next = True
    return p


def add_table(doc: Document, title: str, headers: list[str], rows: list[list[str]], widths: list[float] | None = None):
    add_table_caption(doc, title)
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = False
    if widths:
        for idx, width in enumerate(widths):
            table.columns[idx].width = Cm(width)
    hdr = table.rows[0].cells
    for idx, header in enumerate(headers):
        set_cell_text(hdr[idx], header, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        shade_cell(hdr[idx], "EDEDED")
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            align = WD_ALIGN_PARAGRAPH.CENTER if len(value) < 18 and idx > 0 else WD_ALIGN_PARAGRAPH.LEFT
            set_cell_text(cells[idx], value, align=align)
    paragraph(doc, "", first_line=False)
    return table


def add_listing(doc: Document, title: str, code: str):
    global listing_counter
    listing_counter += 1
    p = paragraph(doc, f"Листинг {listing_counter} - {title}", first_line=False, align=WD_ALIGN_PARAGRAPH.LEFT)
    p.paragraph_format.keep_with_next = True
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    shade_cell(cell, "F5F5F5")
    set_cell_margins(cell, top=140, start=160, bottom=140, end=160)
    cell.text = ""
    for line_no, line in enumerate(code.strip("\n").splitlines(), 1):
        cp = cell.add_paragraph() if line_no > 1 else cell.paragraphs[0]
        cp.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        cp.paragraph_format.space_after = Pt(0)
        cp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = cp.add_run(line)
        run.font.name = "Courier New"
        run.font.size = Pt(9)
        run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Courier New")
    paragraph(doc, "", first_line=False)
    return table


def add_figure(doc: Document, image_path: Path, caption: str, *, width_cm: float = 15.5, page_break: bool = False):
    global figure_counter
    if page_break:
        doc.add_page_break()
    figure_counter += 1
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run()
    run.add_picture(str(image_path), width=Cm(width_cm))
    cap = paragraph(
        doc,
        f"Рисунок {figure_counter} – {caption}",
        first_line=False,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    cap.paragraph_format.space_after = Pt(6)
    for run in cap.runs:
        set_run_font(run, size=12)
    return cap


def add_title_page(doc: Document):
    lines_top = [
        "МИНОБРНАУКИ РОССИИ",
        "Федеральное государственное бюджетное образовательное учреждение",
        "высшего образования",
        "«МИРЭА – Российский технологический университет»",
        "Институт информационных технологий",
        "Кафедра инструментального и прикладного программного обеспечения",
    ]
    for line in lines_top:
        p = paragraph(doc, line, first_line=False, align=WD_ALIGN_PARAGRAPH.CENTER)
        for run in p.runs:
            set_run_font(run, size=14, bold=("МИНОБРНАУКИ" in line or "МИРЭА" in line))

    for _ in range(4):
        paragraph(doc, "", first_line=False)

    p = paragraph(doc, "КУРСОВАЯ РАБОТА", first_line=False, align=WD_ALIGN_PARAGRAPH.CENTER)
    for run in p.runs:
        set_run_font(run, size=16, bold=True)
    p = paragraph(doc, "по дисциплине «Бэкенд-разработка»", first_line=False, align=WD_ALIGN_PARAGRAPH.CENTER)
    for run in p.runs:
        set_run_font(run, size=14)
    p = paragraph(
        doc,
        "на тему: «Серверная часть веб-приложения для хранения и обмена файлами»",
        first_line=False,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    for run in p.runs:
        set_run_font(run, size=14, bold=True)

    for _ in range(4):
        paragraph(doc, "", first_line=False)

    info = [
        "Студент: Жданов Максим Викторович",
        "Группа: ИКБО-11-23",
        "Руководитель: Рачков Андрей Владимирович, ст. преподаватель",
    ]
    for line in info:
        p = paragraph(doc, line, first_line=False, align=WD_ALIGN_PARAGRAPH.RIGHT)
        for run in p.runs:
            set_run_font(run)

    for _ in range(8):
        paragraph(doc, "", first_line=False)

    p = paragraph(doc, "Москва 2026", first_line=False, align=WD_ALIGN_PARAGRAPH.CENTER)
    for run in p.runs:
        set_run_font(run)


def add_front_matter(doc: Document):
    section_heading(doc, "Реферат", page_break=True)
    paragraph(doc, "Отчет 50 с., 29 рис., 10 табл., 5 листинг., 8 источн., 1 прил.")
    paragraph(
        doc,
        "БЭКЕНД-РАЗРАБОТКА, ВЕБ-ПРИЛОЖЕНИЕ, ФАЙЛОВОЕ ХРАНИЛИЩЕ, FASTAPI, "
        "REST API, POSTGRESQL, MINIO, MULTIPART UPLOAD, JWT, NEXT.JS",
    )
    abstract_paragraphs = [
        "Объектом разработки является серверная часть веб-приложения для хранения, версионирования и обмена файлами.",
        "Предметом разработки выступает backend-система, обеспечивающая прием файлов, хранение метаданных, выдачу временных ссылок и разграничение доступа пользователей.",
        "Цель работы - разработать серверную часть веб-приложения «Gosling Drive» с REST API, интеграцией с S3-совместимым объектным хранилищем и базовым пользовательским интерфейсом.",
        "В работе выполнен анализ предметной области, сформированы требования, спроектирована архитектура приложения и база данных, реализованы основные API-сценарии, включая регистрацию, аутентификацию, загрузку файлов, создание новых версий, публичные ссылки и выдачу прав другим пользователям.",
        "Результатом является работающее веб-приложение, развертываемое через Docker Compose. Backend реализован на FastAPI и SQLAlchemy, метаданные хранятся в PostgreSQL, а бинарное содержимое файлов размещается в MinIO. Frontend на Next.js используется для демонстрации ключевых пользовательских операций.",
        "Практическая значимость разработки состоит в создании основы для частного файлового хранилища, которое может применяться в учебных и внутренних корпоративных сценариях, где важны контроль доступа, простота развертывания и независимость от внешних облачных провайдеров.",
    ]
    for text in abstract_paragraphs:
        paragraph(doc, text)

    section_heading(doc, "Термины и определения", page_break=True)
    paragraph(doc, "В настоящем отчете применяют следующие термины с соответствующими определениями.")
    rows = [
        ["Backend", "серверная часть приложения, отвечающая за бизнес-логику, хранение данных и обработку запросов"],
        ["REST API", "программный интерфейс, использующий HTTP-методы и ресурсную модель для обмена данными"],
        ["JWT", "токен, содержащий подписанные сведения о пользователе и применяемый для аутентификации"],
        ["Multipart upload", "механизм загрузки большого файла отдельными частями с последующей сборкой на стороне хранилища"],
        ["S3", "интерфейс объектного хранилища, совместимый с Amazon Simple Storage Service"],
        ["Presigned URL", "временная подписанная ссылка, позволяющая выполнить операцию с объектом без передачи постоянных ключей доступа"],
        ["Версия файла", "зафиксированное состояние содержимого файла, связанное с номером версии и отдельной записью метаданных"],
        ["Публичная ссылка", "токенизированная ссылка для скачивания файла без входа в систему с учетом ограничений срока или числа скачиваний"],
    ]
    add_table(doc, "Термины и определения", ["Термин", "Определение"], rows, widths=[4.3, 11.8])

    section_heading(doc, "Перечень сокращений и обозначений", page_break=True)
    paragraph(doc, "В настоящем отчете применяют следующие сокращения и обозначения.")
    rows = [
        ["API", "Application Programming Interface"],
        ["CRUD", "Create, Read, Update, Delete"],
        ["DB", "Database"],
        ["HTTP", "HyperText Transfer Protocol"],
        ["JSON", "JavaScript Object Notation"],
        ["JWT", "JSON Web Token"],
        ["ORM", "Object-Relational Mapping"],
        ["S3", "Simple Storage Service"],
        ["SQL", "Structured Query Language"],
        ["UI", "User Interface"],
    ]
    add_table(doc, "Сокращения и обозначения", ["Сокращение", "Расшифровка"], rows, widths=[4.0, 12.0])

    section_heading(doc, "Введение", page_break=True)
    intro = [
        "В современных условиях пользователи все чаще работают с цифровыми документами, изображениями, архивами и служебными файлами, которые должны быть доступны с разных устройств и при этом оставаться защищенными. Использование публичных облачных сервисов удобно, но не всегда подходит для учебных и внутренних проектов: данные могут зависеть от сторонней платформы, политика доступа задается внешним сервисом, а возможность доработки логики ограничена.",
        "Разработка собственного файлового хранилища позволяет изучить типовые задачи backend-разработки на практическом примере: проектирование REST API, хранение метаданных, работу с объектным хранилищем, аутентификацию пользователей, контроль доступа, обработку ошибок и тестирование пользовательских сценариев.",
        "Объектом разработки является серверная часть веб-приложения для хранения и обмена файлами. Предметом разработки является программная реализация backend-системы, которая связывает учетные записи пользователей, метаданные файлов, объектное хранилище и веб-интерфейс.",
        "Цель курсовой работы - разработать серверную часть веб-приложения «Gosling Drive», обеспечивающую загрузку файлов, просмотр списка объектов, скачивание, создание новых версий и обмен доступом через публичные ссылки или права для других пользователей.",
        "Для достижения цели необходимо решить следующие задачи: провести анализ предметной области; сформировать функциональные и нефункциональные требования; выбрать технологический стек; спроектировать архитектуру и структуру базы данных; реализовать REST API; интегрировать S3-совместимое хранилище; подготовить веб-интерфейс для демонстрации; провести тестирование и описать запуск системы.",
        "Методами работы являются анализ аналогов, проектирование информационной системы, объектно-реляционное моделирование, реализация серверной логики на Python, интеграционное тестирование и проверка работы приложения в локальном контейнерном окружении.",
        "Основная часть отчета включает анализ предметной области, проектирование системы, описание реализации backend и frontend, тестирование, запуск и эксплуатацию. В заключении приведены результаты работы и направления дальнейшего развития проекта.",
    ]
    for text in intro:
        paragraph(doc, text)


def add_analysis(doc: Document):
    section_heading(doc, "Основная часть", page_break=True)
    section_heading(doc, "1 Анализ предметной области", page_break=False)
    subsection(doc, "1.1 Теоретические аспекты систем хранения файлов")
    texts = [
        "Система хранения файлов представляет собой программный комплекс, который принимает бинарные объекты от пользователя, сохраняет сведения о них и предоставляет контролируемый доступ к скачиванию. В отличие от простого размещения файлов на диске, полноценное приложение должно учитывать владельца, состояние загрузки, версию, тип содержимого, размер, права доступа и историю операций.",
        "В backend-системах задача хранения обычно разделяется на две части. Метаданные, то есть имена файлов, идентификаторы владельцев, размеры, статусы и связи между сущностями, помещаются в реляционную базу данных. Содержимое файлов хранится отдельно в файловой системе или объектном хранилище. Такой подход уменьшает нагрузку на базу данных, упрощает масштабирование и позволяет использовать специализированные механизмы загрузки больших объектов.",
        "Для веб-приложений важен безопасный путь загрузки. Если файл сначала целиком отправляется на backend, сервер становится узким местом по памяти и пропускной способности. Поэтому в проекте применен подход с multipart upload: backend создает сессию, выдает временные ссылки на загрузку частей, браузер передает части напрямую в MinIO, после чего backend завершает сессию и фиксирует метаданные.",
        "Ключевыми свойствами подобных систем являются доступность данных, целостность метаданных, ограничение прав, устойчивость к ошибкам загрузки и возможность восстановления состояния. Для учебного проекта особенно важно, чтобы эти свойства были реализованы прозрачно и могли быть проверены через API, тесты и пользовательский интерфейс.",
    ]
    for text in texts:
        paragraph(doc, text)
    paragraph(doc, "Основные возможности файлового хранилища включают:")
    for item in [
        "регистрацию и вход пользователей;",
        "загрузку файлов и контроль статуса загрузки;",
        "хранение метаданных файлов и версий;",
        "выдачу временных ссылок на скачивание;",
        "создание публичных ссылок с ограничениями;",
        "предоставление прав чтения или записи другим пользователям;",
        "ведение журнала значимых действий.",
    ]:
        bullet(doc, item)

    subsection(doc, "1.2 Анализ существующих решений")
    paragraph(
        doc,
        "На рынке существует множество сервисов хранения файлов, однако их архитектура и возможности различаются. Для курсовой работы важно не только использовать готовый сервис, но и понять, какие функции должны быть реализованы на стороне собственного backend-приложения.",
    )
    add_table(
        doc,
        "Сравнение решений для хранения и обмена файлами",
        ["Критерий", "Google Drive", "Dropbox", "Yandex Disk", "Собственная система"],
        [
            ["Контроль backend-логики", "Низкий", "Низкий", "Низкий", "Полный"],
            ["Изменение модели доступа", "Ограничено", "Ограничено", "Ограничено", "Проектируется под задачу"],
            ["Хранение в локальном контуре", "Нет", "Нет", "Нет", "Возможно"],
            ["Интеграция с учебным API", "Через внешние SDK", "Через внешние SDK", "Через внешние SDK", "Нативная"],
            ["Стоимость масштабирования", "Зависит от тарифа", "Зависит от тарифа", "Зависит от тарифа", "Зависит от инфраструктуры"],
            ["Прозрачность реализации", "Закрытая", "Закрытая", "Закрытая", "Открытая для анализа"],
        ],
        widths=[3.5, 3.0, 3.0, 3.0, 3.6],
    )
    for text in [
        "Готовые облачные сервисы подходят для ежедневной пользовательской работы, но в рамках курсовой работы они не позволяют продемонстрировать проектирование собственного REST API, хранение метаданных и реализацию правил доступа. Поэтому для разработки выбрана собственная система с локальным объектным хранилищем.",
        "Собственная реализация дает возможность управлять схемой данных, форматом API, временем жизни ссылок и политикой версионирования. При этом сохраняется совместимость с распространенным S3-интерфейсом, что упрощает дальнейший перенос на внешнее или промышленное объектное хранилище.",
    ]:
        paragraph(doc, text)

    subsection(doc, "1.3 Формирование требований к системе")
    paragraph(doc, "На основе анализа предметной области сформулированы функциональные требования к приложению.")
    requirements = [
        ["F1", "Регистрация пользователя", "Система должна создавать учетную запись по email, имени и паролю."],
        ["F2", "Аутентификация", "Пользователь должен получать JWT-токен и использовать его для защищенных операций."],
        ["F3", "Загрузка файла", "Система должна создавать upload-сессию и принимать файл multipart-потоком."],
        ["F4", "Список файлов", "Пользователь должен видеть собственные и доступные ему файлы."],
        ["F5", "Скачивание", "Система должна выдавать временную ссылку на скачивание готового файла."],
        ["F6", "Версионирование", "Пользователь должен загружать новую версию существующего файла."],
        ["F7", "Публичная ссылка", "Владелец должен создавать ссылку для скачивания без входа в систему."],
        ["F8", "Разграничение доступа", "Владелец должен выдавать другому пользователю право чтения или записи."],
        ["F9", "Аудит", "Система должна фиксировать значимые действия в журнале."],
    ]
    add_table(doc, "Функциональные требования", ["Код", "Требование", "Описание"], requirements, widths=[1.6, 4.2, 10.3])
    paragraph(doc, "Нефункциональные требования включают:")
    for item in [
        "использование реляционной базы данных для метаданных;",
        "отделение содержимого файлов от метаданных;",
        "работу с S3-совместимым хранилищем;",
        "контейнерный запуск приложения и инфраструктуры;",
        "обработку ошибок с корректными HTTP-кодами;",
        "покрытие ключевых сценариев автоматизированными тестами;",
        "наличие Swagger/OpenAPI-документации для проверки API.",
    ]:
        bullet(doc, item)

    paragraph(
        doc,
        "С точки зрения пользовательских ролей система ориентирована на владельца файла, пользователя с предоставленным доступом и внешнего получателя публичной ссылки. Владелец управляет файлом и ссылками, пользователь с доступом может читать или обновлять файл в зависимости от права, а внешний получатель получает только временную ссылку на скачивание.",
    )


def add_design(doc: Document):
    section_heading(doc, "2 Проектирование системы", page_break=True)
    subsection(doc, "2.1 Архитектура приложения")
    for text in [
        "Приложение спроектировано как клиент-серверная система. Пользователь взаимодействует с Next.js-интерфейсом, frontend обращается к REST API, backend выполняет проверку прав и управляет метаданными, а объектное хранилище MinIO принимает и отдает содержимое файлов через временные ссылки.",
        "В backend-части выделены слои presentation, application, domain и infrastructure. Слой presentation содержит маршрутизаторы FastAPI и зависимости авторизации. Application включает схемы запросов и сервисные функции. Domain содержит перечисления и смысловые сущности. Infrastructure отвечает за базу данных, безопасность и S3-клиент.",
    ]:
        paragraph(doc, text)
    add_figure(doc, ASSETS / "diagram_architecture.png", "Слоистая архитектура приложения Gosling Drive", page_break=False)
    paragraph(
        doc,
        "Такое разделение уменьшает связность: маршрутизаторы не знают детали формирования ключей объектов, а S3-клиент не содержит пользовательской логики. Это упрощает тестирование и позволяет заменить инфраструктурные компоненты без переписывания API-контрактов.",
    )

    subsection(doc, "2.2 Обоснование выбора технологического стека")
    paragraph(
        doc,
        "Основным языком реализации выбран Python, поскольку он широко применяется для backend-разработки, имеет зрелую экосистему веб-фреймворков и хорошо подходит для быстрого создания REST API. Для серверной части выбран FastAPI.",
    )
    add_table(
        doc,
        "Сравнение Python-фреймворков для backend-разработки",
        ["Критерий", "FastAPI", "Django", "Flask"],
        [
            ["Тип", "Современный API-фреймворк", "Полнофункциональный фреймворк", "Микрофреймворк"],
            ["OpenAPI", "Генерируется автоматически", "Требуются расширения", "Требуются расширения"],
            ["Валидация", "Через Pydantic", "Через формы/DRF", "Через сторонние библиотеки"],
            ["Производительность", "Высокая", "Средняя", "Средняя"],
            ["Гибкость архитектуры", "Высокая", "Средняя", "Высокая"],
            ["Подходит для REST API", "Да", "Да, через DRF", "Да"],
        ],
        widths=[3.4, 4.2, 4.2, 4.2],
    )
    for text in [
        "FastAPI выбран из-за встроенной поддержки OpenAPI, удобной интеграции с Pydantic и ясной модели зависимостей. В проекте это позволило быстро описать схемы запросов, проверку токена и типизированные ответы.",
        "Для работы с базой данных использован SQLAlchemy, поскольку он позволяет явно описать модели, связи и ограничения. Миграции выполняются через Alembic, что делает изменение схемы базы данных управляемым и воспроизводимым.",
        "Для хранения метаданных выбран PostgreSQL, так как он поддерживает транзакции, индексы, внешние ключи и ограничения уникальности. Для хранения содержимого файлов выбран MinIO, совместимый с S3 API и удобный для локального контейнерного запуска.",
    ]:
        paragraph(doc, text)
    add_table(
        doc,
        "Выбор средств хранения данных",
        ["Компонент", "Выбранное решение", "Назначение", "Причина выбора"],
        [
            ["Метаданные", "PostgreSQL 16", "Пользователи, файлы, версии, права, аудит", "Транзакционность и строгая схема"],
            ["Содержимое", "MinIO", "Бинарные объекты файлов", "S3-совместимость и presigned URL"],
            ["Миграции", "Alembic", "Изменение структуры БД", "Версионирование схемы"],
            ["Контейнеризация", "Docker Compose", "Запуск зависимостей", "Повторяемость окружения"],
        ],
        widths=[3.0, 3.6, 5.0, 4.5],
    )

    subsection(doc, "2.3 Проектирование базы данных")
    paragraph(
        doc,
        "Схема данных построена вокруг пользователя и файла. Пользователь владеет файлами и папками, файл может иметь несколько версий, загрузка фиксируется отдельной upload-сессией, а доступ и публичные ссылки вынесены в самостоятельные таблицы.",
    )
    add_table(
        doc,
        "Основные сущности базы данных",
        ["Сущность", "Назначение"],
        [
            ["User", "учетная запись пользователя, email, имя, пароль, роль и активность"],
            ["Folder", "папка пользователя с поддержкой родительской папки"],
            ["File", "метаданные текущего состояния файла и ссылка на объект в MinIO"],
            ["FileVersion", "история версий файла с номером версии и объектным ключом"],
            ["UploadSession", "состояние multipart-загрузки, размер части, число частей и статус"],
            ["ShareLink", "публичная ссылка, токен, срок действия и лимит скачиваний"],
            ["AccessGrant", "право чтения или записи для другого пользователя"],
            ["AuditLog", "журнал действий пользователя или внешнего скачивания"],
        ],
        widths=[4.0, 12.1],
    )
    add_figure(doc, ASSETS / "diagram_er.png", "Упрощенная ER-схема базы данных")
    for text in [
        "Для предотвращения дублирования применяются ограничения уникальности: пользователь не может создать две папки с одинаковым именем внутри одного родителя, а также два активных файла с одинаковым именем в одной папке. Версия файла уникальна внутри конкретного файла по номеру версии.",
        "Статусы файлов и upload-сессий позволяют отличать активную загрузку от завершенной или прерванной. Это важно для корректной обработки ошибок: незавершенный объект не должен отображаться пользователю как готовый к скачиванию.",
    ]:
        paragraph(doc, text)
    add_figure(doc, ASSETS / "diagram_storage_split.png", "Разделение метаданных и содержимого файлов")

    subsection(doc, "2.4 Проектирование модели доступа")
    for text in [
        "Модель доступа строится на владельце файла. Если пользователь является владельцем, он может читать и изменять файл. Если пользователь не является владельцем, система ищет запись AccessGrant и разрешает операцию только при наличии нужного права.",
        "Публичная ссылка является отдельным механизмом. Она не раскрывает учетную запись владельца и не требует JWT-токена. Вместо этого проверяется случайный токен, активность ссылки, срок действия и лимит скачиваний.",
    ]:
        paragraph(doc, text)
    add_figure(doc, ASSETS / "diagram_security.png", "Модель доступа к файлу")
    paragraph(
        doc,
        "Такой подход позволяет разделить внутреннюю авторизацию и внешний обмен файлами. Пользователь с JWT работает через защищенные маршруты, а внешний получатель использует только ограниченный публичный маршрут.",
    )

    subsection(doc, "2.5 Проектирование сценария multipart-загрузки")
    for text in [
        "Загрузка файла реализована в несколько этапов. Сначала frontend отправляет на backend имя, размер и тип файла. Backend проверяет права и ограничения, создает запись File со статусом uploading, инициирует multipart upload в MinIO и возвращает параметры сессии.",
        "Далее браузер запрашивает presigned URL для каждой части и отправляет часть напрямую в MinIO методом PUT. После успешной отправки всех частей frontend передает backend список ETag, backend завершает multipart upload, меняет статус файла на ready и создает запись FileVersion.",
    ]:
        paragraph(doc, text)
    add_figure(doc, ASSETS / "diagram_upload_flow.png", "Последовательность multipart-загрузки файла")
    paragraph(
        doc,
        "Главное преимущество этой схемы состоит в том, что backend не проксирует содержимое файла. Он остается ответственным за безопасность, метаданные и завершение операции, но поток данных идет напрямую между браузером и объектным хранилищем.",
    )


def add_backend(doc: Document):
    section_heading(doc, "3 Реализация серверной части", page_break=True)
    subsection(doc, "3.1 Структура проекта")
    for text in [
        "Серверная часть приложения расположена в каталоге app. Точка входа app/main.py создает FastAPI-приложение, подключает CORS, выполняет проверку bucket в MinIO во время startup и включает общий API-router с префиксом /api.",
        "Маршрутизаторы разделены по предметным областям: auth, folders, files, uploads, shares, public, access, audit и storage. Такое разделение делает API удобным для сопровождения: каждый файл отвечает за ограниченный набор операций.",
    ]:
        paragraph(doc, text)
    add_listing(
        doc,
        "Подключение маршрутизаторов FastAPI",
        """
api_router = APIRouter()
api_router.include_router(auth.router)
api_router.include_router(folders.router)
api_router.include_router(files.router)
api_router.include_router(uploads.router)
api_router.include_router(shares.router)
api_router.include_router(public.router)
api_router.include_router(access.router)
api_router.include_router(audit.router)
api_router.include_router(storage.router)
""",
    )
    paragraph(
        doc,
        "Конфигурация приложения вынесена в app/core/config.py и читается из переменных окружения. Это позволяет менять URL базы данных, параметры JWT, адрес MinIO и размер части загрузки без изменения исходного кода.",
    )

    subsection(doc, "3.2 Реализация моделей и миграций")
    for text in [
        "Модели SQLAlchemy описывают таблицы users, folders, files, file_versions, upload_sessions, share_links, access_grants и audit_logs. Для идентификаторов используется UUID в строковом представлении, что упрощает передачу идентификаторов через REST API.",
        "В моделях используются связи relationship, внешние ключи и ограничения UniqueConstraint. Например, FileModel связывается с владельцем, папкой и версиями, а FileVersionModel хранит номер версии, ключ объекта, размер, checksum и ETag.",
        "Миграции Alembic создают начальную структуру базы данных и последующие изменения. При запуске контейнера backend выполняет alembic upgrade head, поэтому база приводится к актуальному состоянию до старта API.",
    ]:
        paragraph(doc, text)

    subsection(doc, "3.3 Реализация REST API")
    paragraph(doc, "Основные конечные точки API сгруппированы по ресурсам.")
    add_table(
        doc,
        "Основные конечные точки REST API",
        ["Группа", "Метод и путь", "Назначение"],
        [
            ["auth", "POST /api/auth/register", "регистрация нового пользователя"],
            ["auth", "POST /api/auth/login", "получение JWT-токена"],
            ["auth", "GET /api/auth/me", "получение данных текущего пользователя"],
            ["folders", "POST /api/folders/", "создание папки"],
            ["files", "GET /api/files/", "список файлов пользователя"],
            ["files", "GET /api/files/{id}/download-url", "выдача временной ссылки на скачивание"],
            ["uploads", "POST /api/uploads/initiate", "создание upload-сессии"],
            ["uploads", "POST /api/uploads/{id}/parts", "выдача ссылок для частей"],
            ["uploads", "POST /api/uploads/{id}/complete", "завершение загрузки"],
            ["shares", "POST /api/shares/", "создание публичной ссылки"],
            ["public", "GET /api/public/{token}", "скачивание по публичному токену"],
            ["access", "POST /api/access/", "выдача доступа другому пользователю"],
            ["storage", "GET /api/storage/stats", "статистика хранилища"],
        ],
        widths=[2.6, 6.2, 7.3],
    )
    for text in [
        "Для входных и выходных данных используются Pydantic-схемы. Они задают обязательные поля, ограничения длины, минимальные значения размеров и формат email. Ошибки валидации автоматически возвращаются клиенту с HTTP-кодом 422.",
        "В маршрутах применяется зависимость get_current_user. Она извлекает bearer-токен, декодирует JWT, находит пользователя в базе данных и отклоняет запрос при отсутствии пользователя или невалидном токене.",
    ]:
        paragraph(doc, text)

    add_figure(doc, SCREEN / "08_swagger_overview.png", "Общая Swagger-документация API")
    add_figure(doc, SCREEN / "09_swagger_auth_folders.png", "Маршруты аутентификации и папок в Swagger")
    add_figure(doc, SCREEN / "10_swagger_files.png", "Маршруты работы с файлами в Swagger")
    add_figure(doc, SCREEN / "11_swagger_uploads.png", "Маршруты upload-сессий в Swagger")
    add_figure(doc, SCREEN / "12_swagger_shares_public.png", "Маршруты публичных ссылок в Swagger")
    add_figure(doc, SCREEN / "13_swagger_access_audit_storage.png", "Маршруты доступа, аудита и статистики в Swagger")

    subsection(doc, "3.4 Реализация S3-хранилища")
    for text in [
        "Интеграция с MinIO реализована в классе S3Storage. Он создает внутренний клиент для операций backend и публичный клиент для генерации ссылок, доступных браузеру. Это разделение нужно потому, что внутри Docker Compose backend обращается к MinIO по имени сервиса, а браузер пользователя - через localhost.",
        "При старте приложение проверяет существование bucket и настраивает CORS, чтобы браузер мог отправлять PUT-запросы по presigned URL. Для загрузки части используется generate_presigned_url с методом upload_part, а для скачивания - get_object.",
    ]:
        paragraph(doc, text)
    add_listing(
        doc,
        "Создание upload-сессии",
        """
upload_session = UploadSessionModel(
    file_id=file.id,
    owner_id=current_user.id,
    provider_upload_id=provider_upload_id,
    bucket=settings.s3_bucket_name,
    object_key=object_key,
    target_version_number=target_version_number,
    target_content_type=content_type,
    target_size_bytes=size_bytes,
    part_size=settings.upload_part_size,
    total_parts=total_parts,
)
""",
    )
    add_figure(doc, SCREEN / "20_minio_login.png", "Страница входа в консоль MinIO")
    add_figure(doc, SCREEN / "21_minio_console.png", "Bucket gosling-drive в консоли MinIO")

    subsection(doc, "3.5 Аудит и обработка ошибок")
    for text in [
        "Для значимых действий вызывается сервис write_audit. В журнал заносятся регистрация, вход, начало и завершение загрузки, создание папки, создание публичной ссылки, выдача доступа и получение ссылок скачивания. Запись содержит пользователя, тип ресурса, идентификатор и дополнительные метаданные.",
        "Ошибки обрабатываются через HTTPException с понятными HTTP-кодами: 401 для невалидного токена, 403 для отключенного пользователя, 404 для недоступного ресурса, 409 для конфликтов имен или состояния файла, 413 для слишком большого файла и 422 для неверных входных данных.",
    ]:
        paragraph(doc, text)
    add_figure(doc, SCREEN / "14_api_health.png", "Проверка health endpoint")
    add_figure(doc, SCREEN / "15_api_me_json.png", "Получение данных текущего пользователя")


def add_frontend(doc: Document):
    section_heading(doc, "4 Реализация веб-интерфейса приложения", page_break=True)
    subsection(doc, "4.1 Общая структура интерфейса")
    for text in [
        "Frontend реализован на Next.js 15 с App Router, TypeScript и Tailwind CSS. Компоненты интерфейса находятся в каталоге frontend/components, а функции обращения к API и загрузки файлов - в frontend/lib.",
        "Главная страница отображает разные состояния: загрузку приложения, гостевой режим и рабочую область авторизованного пользователя. Токен хранится в localStorage под ключом gosling_drive_token и добавляется к запросам в заголовке Authorization.",
    ]:
        paragraph(doc, text)
    add_figure(doc, SCREEN / "01_web_login.png", "Страница входа пользователя")
    add_figure(doc, SCREEN / "02_web_registration.png", "Форма регистрации пользователя")

    subsection(doc, "4.2 Рабочая область пользователя")
    for text in [
        "После входа пользователь видит верхнюю панель с email, числом файлов и общим размером хранилища. Ниже расположены карточки загрузки файла и создания папки, а также таблица файлов с основными действиями.",
        "Таблица файлов показывает имя, размер, номер текущей версии, дату создания, статус и набор команд. Пользователь может скачать файл, создать публичную ссылку или загрузить новую версию.",
    ]:
        paragraph(doc, text)
    add_figure(doc, SCREEN / "03_web_dashboard_files.png", "Рабочая область пользователя с загруженными файлами")
    add_figure(doc, SCREEN / "04_web_upload_area.png", "Карточки загрузки файла и создания папки")
    add_figure(doc, SCREEN / "05_web_file_table.png", "Таблица файлов и доступные действия")
    add_figure(doc, SCREEN / "06_web_mobile_dashboard.png", "Адаптивное отображение интерфейса на мобильной ширине", width_cm=7.0)

    subsection(doc, "4.3 Загрузка файлов через интерфейс")
    for text in [
        "Компонент UploadCard содержит поле выбора файла, кнопку загрузки и индикатор прогресса. При отправке формы вызывается функция uploadLargeFile, которая инициирует сессию, запрашивает ссылки для частей, отправляет данные в MinIO и завершает upload.",
        "В реализации предусмотрены четыре параллельных воркера для загрузки частей. Для небольших файлов используется одна часть, но тот же механизм работает и для крупных объектов, размер которых превышает размер части.",
    ]:
        paragraph(doc, text)
    add_listing(
        doc,
        "Клиентская загрузка файла по частям",
        """
const workerCount = Math.min(PART_WORKER_CONCURRENCY, session.total_parts);
const workers = Array.from({ length: workerCount }, async () => {
  while (nextPart <= session.total_parts) {
    const partNumber = nextPart;
    nextPart += 1;
    const start = (partNumber - 1) * session.part_size;
    const end = Math.min(start + session.part_size, file.size);
    const blob = file.slice(start, end);
    const urlPayload = await api(`/uploads/${session.upload_session_id}/parts`, {
      method: "POST",
      body: JSON.stringify({ part_numbers: [partNumber] }),
    });
    await fetch(urlPayload.urls[0].url, { method: "PUT", body: blob });
  }
});
""",
    )
    add_figure(doc, SCREEN / "22_web_file_selected.png", "Выбор файла перед загрузкой через веб-интерфейс")

    subsection(doc, "4.4 Публичные ссылки")
    paragraph(
        doc,
        "Для обмена файлом пользователь нажимает кнопку «Ссылка». Frontend отправляет запрос POST /api/shares/, получает токен, формирует адрес /share/{token} и копирует его в буфер обмена. Отдельная публичная страница позволяет получателю запросить временную ссылку на скачивание.",
    )
    add_figure(doc, SCREEN / "23_web_share_toast.png", "Создание публичной ссылки из таблицы файлов")
    add_figure(doc, SCREEN / "07_web_public_share.png", "Публичная страница скачивания файла")


def add_testing(doc: Document):
    section_heading(doc, "5 Тестирование и отладка", page_break=True)
    subsection(doc, "5.1 Подход к тестированию")
    for text in [
        "Тестирование выполнялось на двух уровнях. Первый уровень - ручная проверка через браузер, Swagger и реальные HTTP-запросы к локально запущенному приложению. Второй уровень - автоматизированные тесты FastAPI с in-memory базой данных и подмененным S3-хранилищем.",
        "Во время подготовки отчета приложение было запущено в Docker Compose: backend доступен на http://localhost:8000, frontend на http://localhost:3000, MinIO API на http://localhost:9000, консоль MinIO на http://localhost:9001.",
    ]:
        paragraph(doc, text)
    add_table(
        doc,
        "Проверенные сценарии",
        ["Сценарий", "Ожидаемый результат", "Результат"],
        [
            ["Регистрация пользователя", "создается новая учетная запись", "успешно"],
            ["Вход с корректным паролем", "возвращается JWT-токен", "успешно"],
            ["Вход с неверным паролем", "возвращается 401", "успешно"],
            ["Загрузка файла", "файл получает статус ready", "успешно"],
            ["Создание новой версии", "номер версии увеличивается", "успешно"],
            ["Скачивание файла", "возвращается presigned URL", "успешно"],
            ["Публичная ссылка", "внешний маршрут выдает ссылку скачивания", "успешно"],
            ["Доступ другому пользователю", "читатель видит файл с выданным правом", "успешно"],
            ["Защищенный маршрут без токена", "возвращается 401", "успешно"],
        ],
        widths=[5.0, 7.0, 4.0],
    )

    subsection(doc, "5.2 Проверка API в браузере")
    paragraph(
        doc,
        "Для проверки API использовались реальные данные, созданные через HTTP-запросы: демонстрационный пользователь, три файла, новая версия одного файла, публичная ссылка и право чтения для второго пользователя. Ниже приведены снимки JSON-ответов работающего приложения.",
    )
    add_figure(doc, SCREEN / "16_api_files_json.png", "Список файлов пользователя в JSON-ответе")
    add_figure(doc, SCREEN / "17_api_storage_stats_json.png", "Статистика хранилища пользователя")
    add_figure(doc, SCREEN / "18_api_versions_json.png", "Список версий загруженного файла")
    add_figure(doc, SCREEN / "19_api_upload_session_json.png", "Состояние upload-сессии после завершения")

    subsection(doc, "5.3 Автоматизированные тесты")
    for text in [
        "Автоматизированные тесты находятся в каталоге tests. Основной файл test_api_upload_flow.py проверяет полный поток: регистрацию, вход, создание папки, инициирование загрузки, получение ссылок для частей, завершение upload, скачивание, публичную ссылку, создание новой версии, список версий и статистику.",
        "Для изоляции тестов используется SQLite in-memory база данных и FakeStorage. Такой подход позволяет проверить бизнес-логику без реального MinIO и при этом убедиться, что маршруты вызывают ожидаемые методы хранилища.",
        "Дополнительно проверяется хеширование длинных паролей и вспомогательные функции хранилища. Прогон команды uv run pytest завершился результатом: 6 passed, 260 warnings. Предупреждения относятся к deprecated-проверкам в зависимостях FastAPI, Starlette и python-jose при запуске под Python 3.14 и не являются падением тестов.",
    ]:
        paragraph(doc, text)
    add_listing(
        doc,
        "Фрагмент автоматизированного теста upload-flow",
        """
initiate_response = client.post(
    "/api/uploads/initiate",
    headers=auth(token),
    json={
        "filename": "report.txt",
        "size_bytes": 12,
        "content_type": "text/plain",
    },
)
assert initiate_response.status_code == 201
file_payload = complete_one_part_upload(
    client, token, initiate_response.json()["upload_session_id"]
)
assert file_payload["status"] == "ready"
assert file_payload["current_version_number"] == 1
""",
    )


def add_run_ops(doc: Document):
    section_heading(doc, "6 Запуск и эксплуатация системы", page_break=True)
    subsection(doc, "6.1 Контейнерное окружение")
    for text in [
        "Для локального запуска используется Docker Compose. Окружение включает backend, frontend, PostgreSQL, MinIO и одноразовый контейнер minio-init, который создает bucket gosling-drive. Backend зависит от готовности PostgreSQL и завершения инициализации MinIO.",
        "Контейнерный запуск удобен для демонстрации курсовой работы, потому что преподаватель может поднять приложение одной командой и получить одинаковые адреса сервисов.",
    ]:
        paragraph(doc, text)
    add_figure(doc, ASSETS / "diagram_docker.png", "Docker Compose окружение приложения")
    add_listing(
        doc,
        "Команды запуска приложения",
        """
cp .env.example .env
docker-compose up --build

# Backend без Docker:
uv sync --all-groups
uv run alembic upgrade head
uv run uvicorn app.main:app --reload

# Frontend без Docker:
cd frontend
npm install
npm run dev
""",
    )

    subsection(doc, "6.2 Адреса сервисов")
    add_table(
        doc,
        "Локальные адреса приложения",
        ["Сервис", "Адрес", "Назначение"],
        [
            ["Frontend", "http://localhost:3000", "пользовательский веб-интерфейс"],
            ["API", "http://localhost:8000/api", "REST API backend-приложения"],
            ["Swagger", "http://localhost:8000/docs", "интерактивная OpenAPI-документация"],
            ["Healthcheck", "http://localhost:8000/health", "проверка доступности backend"],
            ["MinIO API", "http://localhost:9000", "S3-совместимое объектное хранилище"],
            ["MinIO Console", "http://localhost:9001", "веб-консоль управления MinIO"],
        ],
        widths=[3.5, 5.5, 7.0],
    )
    paragraph(
        doc,
        "При эксплуатации необходимо контролировать значения переменных окружения: JWT-секрет, адрес PostgreSQL, параметры доступа MinIO, имя bucket, срок жизни presigned URL и размер части multipart-загрузки.",
    )

    subsection(doc, "6.3 Рекомендации по дальнейшему развитию")
    paragraph(doc, "Возможные направления развития системы:")
    for item in [
        "добавить полноценную навигацию по папкам во frontend;",
        "реализовать удаление и переименование файлов через интерфейс;",
        "добавить ограничение типов файлов и антивирусную проверку;",
        "реализовать восстановление прерванных upload-сессий;",
        "расширить аудит фильтрацией и отдельной административной страницей;",
        "добавить refresh-токены и управление сессиями пользователя;",
        "подготовить production-конфигурацию с HTTPS и внешним S3-хранилищем.",
    ]:
        bullet(doc, item)


def add_conclusion_sources_appendix(doc: Document):
    section_heading(doc, "Заключение", page_break=True)
    for text in [
        "В ходе выполнения курсовой работы была разработана серверная часть веб-приложения «Gosling Drive», предназначенного для хранения, версионирования и обмена файлами.",
        "Проведен анализ предметной области и существующих решений, сформированы функциональные и нефункциональные требования. На основе требований спроектирована архитектура приложения, модель данных, схема разграничения доступа и сценарий multipart-загрузки.",
        "В рамках реализации создан REST API на FastAPI, описаны SQLAlchemy-модели, подготовлены миграции Alembic, реализована JWT-аутентификация, работа с PostgreSQL и интеграция с MinIO. Для демонстрации разработан веб-интерфейс на Next.js, позволяющий зарегистрироваться, войти, загрузить файл, увидеть список файлов, создать публичную ссылку и загрузить новую версию.",
        "Особое внимание уделено разделению метаданных и содержимого файлов. Метаданные хранятся в PostgreSQL и участвуют в проверке прав, а бинарные объекты помещаются в MinIO. Для передачи содержимого используются presigned URL, что снижает нагрузку на backend и приближает проект к промышленной схеме работы с объектными хранилищами.",
        "Проведенное тестирование подтвердило корректность ключевых сценариев. Автоматизированные тесты покрывают регистрацию, вход, загрузку, версионирование, публичные ссылки, права доступа и статистику. Ручная проверка через интерфейс, Swagger, JSON-ответы API и консоль MinIO подтвердила работоспособность локального развертывания.",
        "Разработанное приложение может использоваться как учебная основа для дальнейшего развития файлового сервиса. В перспективе проект можно расширить навигацией по папкам, управлением жизненным циклом файлов, расширенным аудитом, production-настройками безопасности и интеграцией с внешним S3-провайдером.",
    ]:
        paragraph(doc, text)

    section_heading(doc, "Список использованных источников", page_break=True)
    sources = [
        "FastAPI Documentation [Электронный ресурс]. - Режим доступа: https://fastapi.tiangolo.com/ (дата обращения: 08.05.2026).",
        "SQLAlchemy Documentation [Электронный ресурс]. - Режим доступа: https://docs.sqlalchemy.org/ (дата обращения: 08.05.2026).",
        "Pydantic Documentation [Электронный ресурс]. - Режим доступа: https://docs.pydantic.dev/ (дата обращения: 08.05.2026).",
        "PostgreSQL Documentation [Электронный ресурс]. - Режим доступа: https://www.postgresql.org/docs/ (дата обращения: 08.05.2026).",
        "MinIO Object Store Documentation [Электронный ресурс]. - Режим доступа: https://min.io/docs/minio/ (дата обращения: 08.05.2026).",
        "Boto3 Documentation [Электронный ресурс]. - Режим доступа: https://boto3.amazonaws.com/v1/documentation/api/latest/index.html (дата обращения: 08.05.2026).",
        "Docker Compose Documentation [Электронный ресурс]. - Режим доступа: https://docs.docker.com/compose/ (дата обращения: 08.05.2026).",
        "ГОСТ 7.32-2017. Система стандартов по информации, библиотечному и издательскому делу. Отчет о научно-исследовательской работе. Структура и правила оформления.",
    ]
    for src in sources:
        numbered(doc, src)

    section_heading(doc, "Приложение А", page_break=True)
    paragraph(doc, "Исходный код проекта")
    paragraph(doc, "Репозиторий проекта доступен по адресу:")
    paragraph(doc, "https://github.com/Flict-dev/gosling-drive.git")
    paragraph(doc, "Основные файлы и каталоги проекта:")
    for item in [
        "app/main.py - точка входа FastAPI-приложения;",
        "app/presentation/api/routers - маршрутизаторы REST API;",
        "app/infrastructure/database/models.py - модели SQLAlchemy;",
        "app/infrastructure/storage/s3.py - интеграция с MinIO/S3;",
        "frontend/components - компоненты пользовательского интерфейса;",
        "frontend/lib/upload.ts - клиентский multipart upload;",
        "tests/test_api_upload_flow.py - интеграционные тесты основного сценария;",
        "docker-compose.yml - локальное окружение приложения.",
    ]:
        bullet(doc, item)


def build():
    OUT.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_document(doc)
    add_title_page(doc)
    add_front_matter(doc)
    add_analysis(doc)
    add_design(doc)
    add_backend(doc)
    add_frontend(doc)
    add_testing(doc)
    add_run_ops(doc)
    add_conclusion_sources_appendix(doc)
    doc.core_properties.title = "Курсовая работа: Gosling Drive"
    doc.core_properties.subject = "Серверная часть веб-приложения для хранения и обмена файлами"
    doc.core_properties.author = "Жданов Максим Викторович"
    doc.save(DOCX_PATH)
    enable_update_fields(DOCX_PATH)
    print(DOCX_PATH)
    print(f"figures={figure_counter} tables={table_counter} listings={listing_counter}")


def enable_update_fields(path: Path):
    with ZipFile(path, "r") as src, NamedTemporaryFile(delete=False, suffix=".docx") as tmp_file:
        tmp_path = Path(tmp_file.name)
        with ZipFile(tmp_file, "w", ZIP_DEFLATED) as dst:
            for item in src.infolist():
                data = src.read(item.filename)
                if item.filename == "word/settings.xml" and b"w:updateFields" not in data:
                    marker = b"</w:settings>"
                    data = data.replace(marker, b'<w:updateFields w:val="true"/>' + marker)
                dst.writestr(item, data)
    tmp_path.replace(path)


if __name__ == "__main__":
    build()
